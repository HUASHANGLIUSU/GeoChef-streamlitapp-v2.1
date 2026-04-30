"""
导出工具模块：支持 Excel / CSV / Word / BibTeX / Markdown 五种格式（数据集）。
以及 TXT / Markdown / Word 三种格式（对话记录）。
所有函数接收对应数据，返回 bytes 供 st.download_button 使用。
"""
import io
import re
from datetime import datetime

import pandas as pd

# 内部字段，导出时跳过
_SKIP = {"_sheet", "_text", "_years"}


def _clean_items(items: list[dict]) -> list[dict]:
    """去掉内部字段，返回干净的字典列表。"""
    cleaned = []
    for item in items:
        row = {k: v for k, v in item.items() if k not in _SKIP}
        cleaned.append(row)
    return cleaned


def _to_dataframe(items: list[dict]) -> pd.DataFrame:
    cleaned = _clean_items(items)
    df = pd.DataFrame(cleaned)
    # fillna 先处理真正的 NaN，再统一转字符串，最后清除残余 "nan" 字符串
    df = df.fillna("").astype(str).replace("nan", "").replace("None", "")
    return df


# ── Excel ────────────────────────────────────────────────────────────────────

def to_excel(items: list[dict]) -> bytes:
    df = _to_dataframe(items)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Datasets")
        # 自动调整列宽
        ws = writer.sheets["Datasets"]
        for col_cells in ws.columns:
            max_len = max(
                len(str(cell.value)) if cell.value else 0
                for cell in col_cells
            )
            ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 50)
    return buf.getvalue()


# ── CSV ──────────────────────────────────────────────────────────────────────

def to_csv(items: list[dict]) -> bytes:
    df = _to_dataframe(items)
    return df.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")


# ── Word (.docx) ─────────────────────────────────────────────────────────────

def to_word(items: list[dict]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx")

    doc = Document()

    # 标题
    title = doc.add_heading("RS-VLM Dataset Export", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Exported {len(items)} dataset(s) on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    doc.add_paragraph("")

    cleaned = _clean_items(items)
    for idx, item in enumerate(cleaned, 1):
        name = str(item.get("Name", item.get("name", f"Dataset {idx}"))).strip()
        doc.add_heading(f"{idx}. {name}", level=2)

        for k, v in item.items():
            val = str(v).strip()
            if not val or val == "nan":
                continue
            p = doc.add_paragraph()
            run_key = p.add_run(f"{k}: ")
            run_key.bold = True
            run_key.font.size = Pt(10)
            run_val = p.add_run(val)
            run_val.font.size = Pt(10)

        doc.add_paragraph("")  # 条目间空行

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── BibTeX ───────────────────────────────────────────────────────────────────

def _make_bibtex_key(item: dict, idx: int) -> str:
    name = str(item.get("Name", item.get("name", ""))).strip()
    year = ""
    years = item.get("_years") or set()
    if years:
        year = min(years)
    # 取名称前 12 个字母数字字符
    slug = re.sub(r"[^a-zA-Z0-9]", "", name)[:12]
    return f"{slug}{year}" if slug else f"dataset{idx}"


def to_bibtex(items: list[dict]) -> bytes:
    lines = []
    for idx, item in enumerate(items, 1):
        key = _make_bibtex_key(item, idx)
        name = str(item.get("Name", item.get("name", ""))).strip()
        year_val = ""
        years = item.get("_years") or set()
        if years:
            year_val = min(years)
        publisher = str(item.get("Publisher", item.get("publisher", ""))).strip()
        if publisher == "nan":
            publisher = ""
        note_parts = []
        for k in ["#Samples", "Modality", "Method", "Type"]:
            v = str(item.get(k, "")).strip()
            if v and v != "nan":
                note_parts.append(f"{k}={v}")
        note = "; ".join(note_parts)

        entry = [f"@misc{{{key},"]
        entry.append(f"  title     = {{{name}}},")
        if year_val:
            entry.append(f"  year      = {{{year_val}}},")
        if publisher:
            entry.append(f"  publisher = {{{publisher}}},")
        if note:
            entry.append(f"  note      = {{{note}}},")
        entry.append("}")
        lines.append("\n".join(entry))

    return "\n\n".join(lines).encode("utf-8")


# ── Markdown ─────────────────────────────────────────────────────────────────

def to_markdown(items: list[dict]) -> bytes:
    cleaned = _clean_items(items)
    if not cleaned:
        return b""

    # 收集所有出现过的列（保持顺序）
    all_keys: list[str] = []
    seen_keys: set[str] = set()
    for item in cleaned:
        for k in item:
            if k not in seen_keys:
                all_keys.append(k)
                seen_keys.add(k)

    lines = [
        "# RS-VLM Dataset Export",
        f"\n> Exported {len(items)} dataset(s) on {datetime.now().strftime('%Y-%m-%d %H:%M')}\n",
    ]

    # 表头
    header = "| " + " | ".join(all_keys) + " |"
    sep    = "| " + " | ".join(["---"] * len(all_keys)) + " |"
    lines += [header, sep]

    for item in cleaned:
        cells = []
        for k in all_keys:
            val = str(item.get(k, "")).strip()
            if val == "nan":
                val = ""
            # 转义管道符，避免破坏表格
            val = val.replace("|", "\\|").replace("\n", " ")
            cells.append(val)
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines).encode("utf-8")


# ── 对话记录导出 ──────────────────────────────────────────────────────────────

def _format_role(role: str, lang: str = "cn") -> str:
    if role == "user":
        return "用户" if lang == "cn" else "User"
    return "助手" if lang == "cn" else "Assistant"


def chat_to_txt(history: list[dict], title: str = "", lang: str = "cn") -> bytes:
    """纯文本格式，每条消息用分隔线隔开。"""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = f"{title}\n导出时间：{ts}\n{'='*60}\n\n" if lang == "cn" else \
             f"{title}\nExported at: {ts}\n{'='*60}\n\n"
    lines = [header]
    for msg in history:
        role_label = _format_role(msg["role"], lang)
        lines.append(f"【{role_label}】\n{msg['content']}\n\n{'─'*40}\n")
    return "".join(lines).encode("utf-8")


def chat_to_markdown(history: list[dict], title: str = "", lang: str = "cn") -> bytes:
    """Markdown 格式，用户/助手用不同标题级别区分。"""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [f"# {title}", f"\n> {'导出时间' if lang == 'cn' else 'Exported at'}: {ts}\n"]
    for i, msg in enumerate(history, 1):
        role_label = _format_role(msg["role"], lang)
        icon = "🧑" if msg["role"] == "user" else "🤖"
        lines.append(f"## {icon} {role_label} ({i})\n")
        lines.append(msg["content"])
        lines.append("\n---\n")
    return "\n".join(lines).encode("utf-8")


def chat_to_word(history: list[dict], title: str = "", lang: str = "cn") -> bytes:
    """Word 格式，用户/助手消息用不同样式区分。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx")

    doc = Document()
    heading = doc.add_heading(title or "Chat Export", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sub = doc.add_paragraph(f"{'导出时间' if lang == 'cn' else 'Exported at'}: {ts}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for msg in history:
        role_label = _format_role(msg["role"], lang)
        icon = "🧑 " if msg["role"] == "user" else "🤖 "
        role_para = doc.add_paragraph()
        run = role_para.add_run(f"{icon}{role_label}")
        run.bold = True
        run.font.size = Pt(11)
        # 用户蓝色，助手绿色
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) if msg["role"] == "user" else RGBColor(0x10, 0xB9, 0x81)

        content_para = doc.add_paragraph(msg["content"])
        content_para.style.font.size = Pt(10)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 泄露检测结果导出 ──────────────────────────────────────────────────────────

def _leak_entries_to_items(entries: list[dict], source_label: str = "") -> list[dict]:
    """
    把泄露检测的 entry 列表转为可导出的 item 列表。
    每条 item 在原始数据集字段基础上，额外注入：
      - LeakSource：查询时使用的原始数据集名称
      - IncludedSources：该数据集包含的所有原始数据集（逗号分隔）
    """
    items = []
    for entry in entries:
        item = dict(entry.get("item", {}))
        item["LeakSource"] = source_label
        sources = entry.get("sources", [])
        item["IncludedSources"] = ", ".join(sources) if sources else ""
        items.append(item)
    return items


def leak_to_excel(entries: list[dict], source_label: str = "") -> bytes:
    return to_excel(_leak_entries_to_items(entries, source_label))


def leak_to_csv(entries: list[dict], source_label: str = "") -> bytes:
    return to_csv(_leak_entries_to_items(entries, source_label))


def leak_to_word(entries: list[dict], source_label: str = "") -> bytes:
    return to_word(_leak_entries_to_items(entries, source_label))


def leak_to_bibtex(entries: list[dict], source_label: str = "") -> bytes:
    return to_bibtex(_leak_entries_to_items(entries, source_label))


def leak_to_markdown(entries: list[dict], source_label: str = "") -> bytes:
    return to_markdown(_leak_entries_to_items(entries, source_label))
